#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
授信材料抽取器 —— 把 .docx / .pdf / .xlsx 材料抽成可读文本 + 表格 JSON + 内嵌图片。

用法:
    python3 extract.py 授信报告.docx [审计报告.pdf ...] -o _work/

产出 (每个输入文件一组):
    _work/<stem>.txt          正文纯文本
    _work/<stem>.tables.json  全部表格 (二维数组)
    _work/media/<stem>_NN.png 内嵌图片 (docx)

重点: 授信报告里的「额度测算表 / 股权结构图 / 征信截图 / 裁判文书截图」
常以图片嵌入, 纯文本抽取会丢失。脚本会显式提示图片张数, 必须逐张读图。
"""
import argparse
import json
import os
import re
import shutil
import subprocess
import sys
import zipfile

IMG_EXT = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".emf", ".wmf", ".tif", ".tiff"}


def log(msg):
    print(msg, flush=True)


# ---------------------------------------------------------------- docx
def extract_docx(path, outdir):
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    stem = os.path.splitext(os.path.basename(path))[0]
    doc = Document(path)

    def iter_block_items(parent):
        from docx.oxml.ns import qn
        body = parent.element.body
        for child in body.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, parent)
            elif child.tag == qn("w:tbl"):
                yield Table(child, parent)

    lines, tables = [], []
    t_idx = 0
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            txt = block.text.strip()
            if txt:
                style = (block.style.name or "").lower()
                lines.append(("#" * 2 + " " + txt) if "heading" in style else txt)
            elif block._p.xpath(".//pic:pic"):
                lines.append("〔此处为内嵌图片，见 media/ 目录，须逐张读图〕")
        else:
            t_idx += 1
            rows = []
            for r in block.rows:
                cells, seen = [], set()
                for c in r.cells:
                    if id(c._tc) in seen:
                        continue
                    seen.add(id(c._tc))
                    cells.append(re.sub(r"\s+", " ", c.text).strip())
                rows.append(cells)
            tables.append({"序号": t_idx, "行数": len(rows), "数据": rows})
            lines.append(f"〔表 {t_idx}：{len(rows)} 行，详见 {stem}.tables.json〕")

    # 内嵌图片
    mediadir = os.path.join(outdir, "media")
    os.makedirs(mediadir, exist_ok=True)
    n_img = 0
    with zipfile.ZipFile(path) as z:
        for name in z.namelist():
            if name.startswith("word/media/") and os.path.splitext(name)[1].lower() in IMG_EXT:
                n_img += 1
                dst = os.path.join(mediadir, f"{stem}_{n_img:02d}{os.path.splitext(name)[1].lower()}")
                with z.open(name) as src, open(dst, "wb") as f:
                    shutil.copyfileobj(src, f)
    return "\n\n".join(lines), tables, n_img


# ---------------------------------------------------------------- pdf
def extract_pdf(path, outdir):
    for cmd in (["pdftotext", "-layout", path, "-"],):
        try:
            out = subprocess.run(cmd, capture_output=True, timeout=180)
            if out.returncode == 0 and out.stdout.strip():
                return out.stdout.decode("utf-8", "replace"), [], 0
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
    try:
        import pdfplumber
        txt, tables = [], []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                txt.append(f"—— 第 {i} 页 ——\n" + (page.extract_text() or ""))
                for t in page.extract_tables() or []:
                    tables.append({"序号": len(tables) + 1, "页": i,
                                   "行数": len(t),
                                   "数据": [[(c or "").strip() for c in row] for row in t]})
        return "\n\n".join(txt), tables, 0
    except ImportError:
        return "〔无法抽取 PDF：缺少 pdftotext 与 pdfplumber，请改用 Read 工具直接读图〕", [], 0


# ---------------------------------------------------------------- xlsx
def extract_xlsx(path, outdir):
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    tables, lines = [], []
    for ws in wb.worksheets:
        rows = [[("" if c is None else str(c)).strip() for c in r]
                for r in ws.iter_rows(values_only=True)]
        rows = [r for r in rows if any(r)]
        tables.append({"序号": len(tables) + 1, "工作表": ws.title,
                       "行数": len(rows), "数据": rows})
        lines.append(f"〔工作表「{ws.title}」：{len(rows)} 行〕")
    return "\n".join(lines), tables, 0


HANDLERS = {".docx": extract_docx, ".docm": extract_docx,
            ".pdf": extract_pdf,
            ".xlsx": extract_xlsx, ".xlsm": extract_xlsx}


def main():
    ap = argparse.ArgumentParser(description="授信材料抽取器")
    ap.add_argument("files", nargs="+")
    ap.add_argument("-o", "--outdir", default="_work")
    args = ap.parse_args()
    os.makedirs(args.outdir, exist_ok=True)

    total_img = 0
    for path in args.files:
        if not os.path.exists(path):
            log(f"✗ 文件不存在：{path}")
            continue
        ext = os.path.splitext(path)[1].lower()
        handler = HANDLERS.get(ext)
        if not handler:
            log(f"— 跳过（不支持 {ext}）：{os.path.basename(path)}")
            continue
        stem = os.path.splitext(os.path.basename(path))[0]
        try:
            text, tables, n_img = handler(path, args.outdir)
        except Exception as e:
            log(f"✗ 抽取失败 {os.path.basename(path)}: {e}")
            continue

        with open(os.path.join(args.outdir, stem + ".txt"), "w", encoding="utf-8") as f:
            f.write(text)
        if tables:
            with open(os.path.join(args.outdir, stem + ".tables.json"), "w", encoding="utf-8") as f:
                json.dump(tables, f, ensure_ascii=False, indent=2)

        total_img += n_img
        log(f"✓ {os.path.basename(path)}  →  正文 {len(text):,} 字符 / 表格 {len(tables)} 个 / 图片 {n_img} 张")

    if total_img:
        log("")
        log(f"⚠ 共检出 {total_img} 张内嵌图片，已存入 {os.path.join(args.outdir, 'media')}/")
        log("  额度测算表、股权结构图、征信与裁判文书截图常以图片嵌入。")
        log("  请逐张 Read 读图取数；读不出的一律计入「待补」，不得默认通过。")


if __name__ == "__main__":
    sys.exit(main())
