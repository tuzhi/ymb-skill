#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
四交付物生成器 —— 由 review.json 一次生成:
  1. 《XXX-审查报告》.docx        八段式正式报告 + 关键复算过程附录
  2. 《XXX-审查点检表》.xlsx      六大类逐条留痕 + 类别汇总 + 穿透抽查
  3. 《XXX-补充资料清单》.md      不符合/待补项汇总, 直接发业务团队
  4. 《XXX-额度测算底稿》.xlsx    营运资金缺口法 + 订单模式 + 担保能力 + 财务勾稽
                                 (额度与担保能力为活公式, 现场可改参数重算)

用法:
    python3 build_outputs.py review.json -o 输出目录/

数据契约见 references/05-review-json数据契约.md
"""
import argparse
import json
import os
import sys
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import openpyxl
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

# ------------------------------------------------------------------ 常量
FILL = {
    "符合":   PatternFill("solid", fgColor="DDF3E0"),
    "不符合": PatternFill("solid", fgColor="FBD5D5"),
    "待补":   PatternFill("solid", fgColor="FDF0C9"),
    "不适用": PatternFill("solid", fgColor="EFEFEF"),
}
HEAD_FILL = PatternFill("solid", fgColor="1F3864")
HEAD_FONT = Font(name="微软雅黑", size=10, bold=True, color="FFFFFF")
BODY_FONT = Font(name="微软雅黑", size=10)
TITLE_FONT = Font(name="微软雅黑", size=13, bold=True, color="1F3864")
THIN = Side(style="thin", color="BFBFBF")
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
WRAP = Alignment(vertical="top", wrap_text=True)
CENTER = Alignment(horizontal="center", vertical="center", wrap_text=True)


def num(v, default=0.0):
    """None / '' / 非数值 → default"""
    try:
        if v is None or v == "":
            return default
        return float(v)
    except (TypeError, ValueError):
        return default


def has(v):
    return v is not None and v != ""


def fmt(v, nd=2):
    if not isinstance(v, (int, float)):
        return str(v)
    return f"{v:,.{nd}f}"


# ================================================================= 复算
def calc_limit(d):
    """营运资金缺口法复算。返回 dict, 含 ok / 缺失参数 / 中间量。"""
    if not d:
        return None
    r = {"缺失参数": []}
    for k in ("上年度销售收入", "上年度净利润", "营业成本"):
        if not has(d.get(k)):
            r["缺失参数"].append(k)
    rev = num(d.get("上年度销售收入"))
    profit = num(d.get("上年度净利润"))
    cost = num(d.get("营业成本"))

    def avg(name):
        b, e = d.get(f"{name}_期初"), d.get(f"{name}_期末")
        if not has(b) and not has(e):
            r["缺失参数"].append(name)
            return 0.0
        if not has(b):
            return num(e)
        if not has(e):
            return num(b)
        return (num(b) + num(e)) / 2

    inv, ar, pre_pay, ap, pre_rec = (avg("存货"), avg("应收账款"), avg("预付账款"),
                                     avg("应付账款"), avg("预收账款"))
    r["平均存货"], r["平均应收账款"] = inv, ar
    r["平均预付账款"], r["平均应付账款"], r["平均预收账款"] = pre_pay, ap, pre_rec

    r["存货周转天数"] = 360 * inv / cost if cost else None
    r["应收账款周转天数"] = 360 * ar / rev if rev else None
    r["预付账款周转天数"] = 360 * pre_pay / cost if cost else None
    r["应付账款周转天数"] = 360 * ap / cost if cost else None
    r["预收账款周转天数"] = 360 * pre_rec / rev if rev else None

    days = [r["存货周转天数"], r["应收账款周转天数"], r["预付账款周转天数"],
            r["应付账款周转天数"], r["预收账款周转天数"]]
    if any(x is None for x in days):
        r["ok"] = False
        r["失效原因"] = "收入或成本为零/缺失，周转天数无法计算"
        return r
    r["周转天数合计"] = days[0] + days[1] + days[2] - days[3] - days[4]
    if r["周转天数合计"] <= 0:
        r["ok"] = False
        r["失效原因"] = ("周转天数合计 ≤ 0，说明占用上游资金大于自身垫资，"
                         "理论上无营运资金缺口 —— 营运资金缺口法失效，建议转订单模式")
        return r

    r["营运资金周转次数"] = 360 / r["周转天数合计"]
    # 销售利润率有两种口径，必须与业务团队一致，否则复算差异全部来自口径而非数据
    caliber = d.get("销售利润率口径") or "净利率"
    if caliber == "毛利率":
        r["销售利润率"] = (rev - cost) / rev if rev else 0.0
    else:
        r["销售利润率"] = profit / rev if rev else 0.0
    r["销售利润率口径"] = caliber
    r["预计销售收入年增长率"] = num(d.get("预计销售收入年增长率"))
    r["营运资金量"] = (rev * (1 - r["销售利润率"]) * (1 + r["预计销售收入年增长率"])
                       / r["营运资金周转次数"])
    r["复算缺口"] = (r["营运资金量"] - num(d.get("自有资金"))
                     - num(d.get("现有流动资金贷款")) - num(d.get("其他渠道营运资金")))
    r["ok"] = True

    # 关键周转参数缺失时，复算结果不可靠 —— 只能判"待补"，不得据此判"不符合"
    r["可靠"] = not r["缺失参数"]

    rep = d.get("报告测算缺口")
    if has(rep):
        r["报告测算缺口"] = num(rep)
        r["差异"] = r["复算缺口"] - r["报告测算缺口"]
        r["差异率"] = r["差异"] / r["报告测算缺口"] if r["报告测算缺口"] else None
        if not r["可靠"]:
            r["判定"] = ("待补 —— 缺失参数（" + "、".join(r["缺失参数"]) +
                        "），复算结果按 0 计入，必然高估营运资金量，仅供参考；"
                        "须补明细账龄表后重算")
        elif r["差异率"] is None:
            r["判定"] = "待补"
        elif abs(r["差异率"]) <= 0.10:
            r["判定"] = "符合（差异 ≤ 10%）"
        else:
            r["判定"] = "不符合（差异 > 10%，须并列两套算式并说明差异来源）"
    apply_amt = d.get("申请额度")
    if has(apply_amt):
        r["申请额度"] = num(apply_amt)
        r["路径A_按缺口核定"] = min(r["申请额度"], max(r["复算缺口"], 0))
        if not r["可靠"]:
            r["额度充足性"] = "参数缺失，复算缺口不具结论意义，额度充足性待重算后判定"
        elif r["复算缺口"] >= r["申请额度"]:
            r["额度充足性"] = "复算缺口已覆盖申请额度"
        else:
            r["额度充足性"] = "复算缺口低于申请额度 —— 提示两条路径：A 按缺口核定 / B 转订单模式"
    return r


def calc_order(d):
    if not d or not d.get("启用"):
        return None
    orders = d.get("订单") or []
    rows, total = [], 0.0
    for i, o in enumerate(orders, 1):
        amt, ratio = num(o.get("金额")), num(o.get("融资比例"), 0.7)
        fin = amt * ratio
        total += fin
        rows.append({"序号": i, "供应商": o.get("供应商", ""), "标的": o.get("标的", ""),
                     "金额": amt, "融资比例": ratio, "可融资额": fin})
    days = sum(num(d.get(k)) for k in ("采购交付天数", "施工验收天数",
                                       "结算天数", "账期天数", "缓冲天数"))
    months = int(-(-days // 30)) if days else 0
    return {"订单明细": rows, "可融资合计": total, "回款周期天数": days,
            "建议期限月": months,
            "期限约束": "期限不得超过单笔订单的实际回款周期；多笔订单按加权平均确定并以最长单笔封顶"}


def calc_guarantee(d):
    if not d:
        return None
    cap = (num(d.get("净资产")) - num(d.get("已担保金额"))
           - num(d.get("或有负债")) - num(d.get("受限资产")))
    amt = num(d.get("本次授信金额"))
    return {**d, "担保能力": cap, "覆盖倍数": (cap / amt if amt else None),
            "判定": ("符合 —— 担保能力可覆盖本次授信金额" if amt and cap >= amt
                    else "不符合 —— 担保能力不足，须追加担保或压降额度")}


def calc_ties(fin, gap=None):
    """财务勾稽 K01–K10 + E05/E06/E10。返回按年份的检查行。"""
    if not fin:
        return None
    seg = {}
    for s in fin.get("分部表") or []:
        y = str(s.get("年份"))
        seg.setdefault(y, {"收入": 0.0, "成本": 0.0})
        seg[y]["收入"] += num(s.get("收入"))
        seg[y]["成本"] += num(s.get("成本"))

    years = fin.get("年度") or []
    last_year = str(years[-1].get("年份")) if years else None
    out = []
    for y in years:
        year = str(y.get("年份"))
        rec = {"年份": year, "检查": []}

        def chk(code, name, calc, stated, tol, unit=""):
            if calc is None:
                rec["检查"].append({"编号": code, "项目": name, "计算值": "—",
                                    "报告值": fmt(stated) if has(stated) else "—",
                                    "差异": "—", "判定": "待补 —— 参数缺失"})
                return
            if not has(stated):
                rec["检查"].append({"编号": code, "项目": name, "计算值": fmt(calc),
                                    "报告值": "—", "差异": "—", "判定": "报告未表述"})
                return
            diff = calc - num(stated)
            rec["检查"].append({"编号": code, "项目": name, "计算值": fmt(calc),
                                "报告值": fmt(num(stated)), "差异": fmt(diff),
                                "判定": "一致" if abs(diff) <= tol else f"不一致（阈值 {tol}{unit}）"})

        ta, tl, eq = y.get("资产合计"), y.get("负债合计"), y.get("所有者权益")
        if has(ta) and has(tl) and has(eq):
            gapv = num(ta) - num(tl) - num(eq)
            rec["检查"].append({"编号": "K01", "项目": "资产 = 负债 + 所有者权益",
                                "计算值": fmt(num(tl) + num(eq)), "报告值": fmt(num(ta)),
                                "差异": fmt(gapv),
                                "判定": "平衡" if abs(gapv) <= 1 else "不平衡（差异 > 1 万元）"})
        chk("K03", "资产负债率 %",
            (num(tl) / num(ta) * 100 if has(ta) and num(ta) else None),
            y.get("文字_资产负债率"), 0.5, "pct")
        ca, cl = y.get("流动资产"), y.get("流动负债")
        chk("K04", "流动比率 %",
            (num(ca) / num(cl) * 100 if has(ca) and has(cl) and num(cl) else None),
            y.get("文字_流动比率"), 5, "pct")
        chk("K05", "速动比率 %",
            ((num(ca) - num(y.get("存货"))) / num(cl) * 100
             if has(ca) and has(cl) and num(cl) else None),
            y.get("文字_速动比率"), 5, "pct")
        rev, cost = y.get("营业收入"), y.get("营业成本")
        gm = ((num(rev) - num(cost)) / num(rev) * 100
              if has(rev) and has(cost) and num(rev) else None)
        chk("K06", "毛利率 %", gm, y.get("文字_毛利率"), 0.5, "pct")
        if year in seg:
            chk("K07", "分部收入合计 vs 营业收入", seg[year]["收入"], rev, max(num(rev) * 0.01, 1))
            chk("K07b", "分部成本合计 vs 营业成本", seg[year]["成本"], cost, max(num(cost) * 0.01, 1))
            sgm = ((seg[year]["收入"] - seg[year]["成本"]) / seg[year]["收入"] * 100
                   if seg[year]["收入"] else None)
            chk("K08", "分部表反算毛利率 %", sgm, y.get("文字_毛利率"), 0.5, "pct")
        if has(y.get("净利润")) and has(rev) and num(rev):
            rec["检查"].append({"编号": "K09", "项目": "净利率 %",
                                "计算值": fmt(num(y["净利润"]) / num(rev) * 100),
                                "报告值": "—", "差异": "—", "判定": "供额度测算取数"})
        if has(y.get("货币资金")) and gap and year == last_year:
            ratio = num(y["货币资金"]) / gap if gap else None
            rec["检查"].append({"编号": "E05", "项目": "货币资金 ÷ 测算资金缺口（倍）",
                                "计算值": fmt(ratio), "报告值": "—", "差异": "—",
                                "判定": ("融资必要性存疑（> 3 倍），须说明资金受限情况"
                                        if ratio and ratio > 3 else
                                        "关注（1–3 倍）" if ratio and ratio > 1 else "正常")})
        if has(y.get("其他应收款")) and has(eq) and num(eq):
            ratio = num(y["其他应收款"]) / num(eq) * 100
            rec["检查"].append({"编号": "E06", "项目": "其他应收款 ÷ 净资产 %",
                                "计算值": fmt(ratio), "报告值": "—", "差异": "—",
                                "判定": ("关联方占款偏大（> 30%），须在风险分析中单列"
                                        if ratio > 30 else "正常")})
        if has(y.get("政府补助")) and has(y.get("净利润")) and num(y["净利润"]):
            ratio = num(y["政府补助"]) / num(y["净利润"]) * 100
            rec["检查"].append({"编号": "E10", "项目": "政府补助 ÷ 净利润 %",
                                "计算值": fmt(ratio), "报告值": "—", "差异": "—",
                                "判定": ("扣除补助后经营性亏损，第一还款来源不成立"
                                        if ratio > 100 else
                                        "经营性还款来源弱，须在风险分析中写明"
                                        if ratio > 50 else "正常")})
        out.append(rec)
    return out


# ================================================================= docx
def set_font(run, cn="仿宋_GB2312", en="Times New Roman", size=14, bold=False):
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn)


def para(doc, text, cn="仿宋_GB2312", size=14, bold=False, align=None,
         first_indent=True, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    if align is not None:
        p.alignment = align
    set_font(p.add_run(text), cn=cn, size=size, bold=bold)
    return p


def add_table(doc, spec):
    data = spec.get("数据") or []
    header = spec.get("表头") or []
    if not data and not header:
        return
    ncol = len(header) if header else max(len(r) for r in data)
    if spec.get("标题"):
        para(doc, spec["标题"], cn="黑体", size=10.5, first_indent=False, space_after=2)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if header:
        cells = t.add_row().cells
        for i, h in enumerate(header[:ncol]):
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(cells[i].paragraphs[0].add_run(str(h)), cn="黑体", size=10.5, bold=True)
    for row in data:
        cells = t.add_row().cells
        for i in range(ncol):
            v = row[i] if i < len(row) else ""
            set_font(cells[i].paragraphs[0].add_run("" if v is None else str(v)),
                     cn="仿宋_GB2312", size=10.5)
    doc.add_paragraph()


def render_sections(doc, sections, level=1):
    for sec in sections or []:
        if sec.get("标题"):
            size = {1: 15, 2: 14, 3: 14}.get(level, 14)
            para(doc, sec["标题"], cn="黑体", size=size, bold=(level == 1),
                 first_indent=(level > 1), space_after=4)
        for p in sec.get("段落") or []:
            para(doc, p)
        for tb in ([sec["表格"]] if isinstance(sec.get("表格"), dict) else sec.get("表格") or []):
            add_table(doc, tb)
        render_sections(doc, sec.get("子节"), level + 1)


def build_docx(data, calcs, path):
    meta, rep = data.get("meta", {}), data.get("报告", {})
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Cm(2.8)
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(14)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")

    title = rep.get("标题") or f"{meta.get('客户名称','')}综合授信{fmt(num(meta.get('授信金额')),0)}万元的审查报告"
    para(doc, title, cn="黑体", size=18, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, space_after=14)

    if rep.get("履职声明"):
        para(doc, "履职声明：" + rep["履职声明"])
    for line in rep.get("抬头") or []:
        para(doc, line)
    doc.add_paragraph()

    render_sections(doc, rep.get("章节"))

    para(doc, rep.get("结尾", "妥否，请审议。"))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(p.add_run(f"审查员：{meta.get('审查员','')}　　　{meta.get('审查日期', date.today().isoformat())}"))

    # ---- 附录：关键复算过程（数字留痕，贷审会可现场复算）
    doc.add_page_break()
    para(doc, "附：关键复算过程", cn="黑体", size=15, bold=True, first_indent=False)

    lim = calcs.get("额度")
    if lim:
        para(doc, "一、授信额度复算（营运资金缺口法）", cn="黑体", size=14, first_indent=False)
        if lim.get("ok"):
            rows = [
                ["存货周转天数", f'360 × {fmt(lim["平均存货"])} ÷ 营业成本', fmt(lim["存货周转天数"])],
                ["应收账款周转天数", f'360 × {fmt(lim["平均应收账款"])} ÷ 营业收入', fmt(lim["应收账款周转天数"])],
                ["预付账款周转天数", f'360 × {fmt(lim["平均预付账款"])} ÷ 营业成本', fmt(lim["预付账款周转天数"])],
                ["应付账款周转天数", f'360 × {fmt(lim["平均应付账款"])} ÷ 营业成本', fmt(lim["应付账款周转天数"])],
                ["预收账款周转天数", f'360 × {fmt(lim["平均预收账款"])} ÷ 营业收入', fmt(lim["预收账款周转天数"])],
                ["周转天数合计", "存货 + 应收 + 预付 − 应付 − 预收", fmt(lim["周转天数合计"])],
                ["营运资金周转次数", "360 ÷ 周转天数合计", fmt(lim["营运资金周转次数"])],
                [f'销售利润率（{lim["销售利润率口径"]}口径）',
                 "净利润 ÷ 营业收入" if lim["销售利润率口径"] == "净利率" else "(营业收入 − 营业成本) ÷ 营业收入",
                 f'{lim["销售利润率"]*100:.2f}%'],
                ["预计销售收入年增长率", "业务团队假设", f'{lim["预计销售收入年增长率"]*100:.2f}%'],
                ["营运资金量", "销售收入 × (1 − 销售利润率) × (1 + 预计增长率) ÷ 周转次数", fmt(lim["营运资金量"])],
                ["复算新增流贷额度", "营运资金量 − 自有资金 − 现有流贷 − 其他渠道", fmt(lim["复算缺口"])],
            ]
            if "报告测算缺口" in lim:
                rows += [["报告测算缺口", "业务团队测算值", fmt(lim["报告测算缺口"])],
                         ["差异 / 差异率", "复算 − 报告",
                          f'{fmt(lim["差异"])}　/　{lim["差异率"]*100:.2f}%' if lim.get("差异率") is not None else fmt(lim["差异"])],
                         ["判定", "阈值 ±10%", lim.get("判定", "")]]
            if "申请额度" in lim:
                rows += [["申请额度", "", fmt(lim["申请额度"])],
                         ["额度充足性", "", lim.get("额度充足性", "")],
                         ["路径 A：按缺口核定", "min(申请额度, 复算缺口)", fmt(lim["路径A_按缺口核定"])]]
            add_table(doc, {"表头": ["项目", "算式", "结果（万元 / %）"], "数据": rows})
            if lim.get("缺失参数"):
                para(doc, f'※ 复算可靠性提示：以下参数缺失，已按 0 计入 —— '
                          f'{"、".join(lim["缺失参数"])}。缺失应付账款、预收账款会**低估**周转天数扣减项，'
                          f'从而**高估**营运资金量与资金缺口，本次复算结果仅供参考，'
                          f'须补充明细账龄表后重算，相应点检项判「待补」。')
        else:
            para(doc, f'营运资金缺口法测算失效：{lim.get("失效原因","参数缺失")}。'
                      f'缺失参数：{"、".join(lim.get("缺失参数") or []) or "无"}。'
                      f'按 references/03 建议转订单模式或要求补充明细账龄表。')

    od = calcs.get("订单")
    if od:
        para(doc, "二、订单模式测算", cn="黑体", size=14, first_indent=False)
        rows = [[o["序号"], o["供应商"], o["标的"], fmt(o["金额"]),
                 f'{o["融资比例"]*100:.0f}%', fmt(o["可融资额"])] for o in od["订单明细"]]
        rows.append(["合计", "", "", "", "", fmt(od["可融资合计"])])
        add_table(doc, {"表头": ["序号", "供应商", "标的", "合同金额（万元）", "融资比例", "可融资额（万元）"],
                        "数据": rows})
        para(doc, f'回款周期倒推：采购交付 + 施工验收 + 结算 + 账期 + 缓冲 = '
                  f'{fmt(od["回款周期天数"],0)} 天，建议授信期限 {od["建议期限月"]} 个月。'
                  f'{od["期限约束"]}。')

    g = calcs.get("担保")
    if g:
        para(doc, "三、担保能力复算", cn="黑体", size=14, first_indent=False)
        para(doc, f'担保能力 = 净资产 {fmt(num(g.get("净资产")))} − 已担保 {fmt(num(g.get("已担保金额")))}'
                  f'（{g.get("已担保口径","口径未注明")}） − 或有负债 {fmt(num(g.get("或有负债")))}'
                  f' − 受限资产 {fmt(num(g.get("受限资产")))} = {fmt(g["担保能力"])} 万元。')
        para(doc, f'本次授信金额 {fmt(num(g.get("本次授信金额")))} 万元，覆盖倍数 '
                  f'{fmt(g["覆盖倍数"]) if g.get("覆盖倍数") is not None else "—"} 倍。{g["判定"]}。')
        para(doc, f'口径与基准日：{g.get("净资产口径","未注明")}；{g.get("基准日","未注明")}。')
        if has(g.get("关注类担保余额")):
            para(doc, f'其中关注类担保余额 {fmt(num(g["关注类担保余额"]))} 万元，须单独评估代偿风险。')

    ties = calcs.get("勾稽")
    if ties:
        para(doc, "四、财务勾稽校验", cn="黑体", size=14, first_indent=False)
        for rec in ties:
            para(doc, f'{rec["年份"]} 年度：', cn="黑体", size=13, first_indent=False, space_after=2)
            add_table(doc, {"表头": ["编号", "校验项", "计算值", "报告值", "差异", "判定"],
                            "数据": [[c["编号"], c["项目"], c["计算值"], c["报告值"],
                                     c["差异"], c["判定"]] for c in rec["检查"]]})

    doc.save(path)
    return path


# ================================================================= xlsx 点检表
def style_sheet(ws, widths, header_row=1):
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    for c in ws[header_row]:
        c.fill, c.font, c.alignment, c.border = HEAD_FILL, HEAD_FONT, CENTER, BORDER
    ws.freeze_panes = ws.cell(row=header_row + 1, column=1)


def build_checklist(data, path):
    wb = openpyxl.Workbook()
    items = data.get("点检表") or []

    ws = wb.active
    ws.title = "点检明细"
    cols = ["编号", "类别", "点检项", "结论", "证据出处", "说明"]
    ws.append(cols)
    for it in items:
        ws.append([it.get(c, "") for c in cols])
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, max_col=len(cols)):
        for c in row:
            c.font, c.alignment, c.border = BODY_FONT, WRAP, BORDER
        concl = row[3].value
        if concl in FILL:
            row[3].fill = FILL[concl]
            row[3].alignment = CENTER
            row[3].font = Font(name="微软雅黑", size=10, bold=(concl == "不符合"))
    style_sheet(ws, [8, 20, 30, 10, 34, 60])

    # 汇总
    ws2 = wb.create_sheet("类别汇总")
    ws2.append(["类别", "符合", "不符合", "待补", "不适用", "合计"])
    cats = []
    for it in items:
        if it.get("类别") not in cats:
            cats.append(it.get("类别"))
    for cat in cats:
        sub = [i for i in items if i.get("类别") == cat]
        ws2.append([cat] + [sum(1 for i in sub if i.get("结论") == k)
                            for k in ("符合", "不符合", "待补", "不适用")] + [len(sub)])
    ws2.append(["总计"] + [sum(1 for i in items if i.get("结论") == k)
                          for k in ("符合", "不符合", "待补", "不适用")] + [len(items)])
    for row in ws2.iter_rows(min_row=2, max_row=ws2.max_row, max_col=6):
        for c in row:
            c.font, c.border = BODY_FONT, BORDER
            if c.column > 1:
                c.alignment = CENTER
    ws2.cell(row=ws2.max_row, column=1).font = Font(name="微软雅黑", size=10, bold=True)
    style_sheet(ws2, [24, 10, 10, 10, 10, 10])

    # 穿透抽查
    ws3 = wb.create_sheet("穿透抽查")
    ws3.append(["抽查项", "做法", "发现"])
    for p in data.get("穿透抽查") or []:
        ws3.append([p.get("抽查项", ""), p.get("做法", ""), p.get("发现", "")])
    if ws3.max_row == 1:
        ws3.append(["⚠ 未记录穿透抽查", "反套路化要求每案至少 2 项", "返回阶段 3 补做"])
    for row in ws3.iter_rows(min_row=2, max_row=ws3.max_row, max_col=3):
        for c in row:
            c.font, c.alignment, c.border = BODY_FONT, WRAP, BORDER
    style_sheet(ws3, [28, 52, 60])

    # 文本一致性
    tc = data.get("文本一致性") or []
    if tc:
        ws4 = wb.create_sheet("文本一致性")
        ws4.append(["指标", "出处1", "出处2", "结论"])
        for t in tc:
            ws4.append([t.get("指标", ""), t.get("出处1", ""), t.get("出处2", ""), t.get("结论", "")])
        for row in ws4.iter_rows(min_row=2, max_row=ws4.max_row, max_col=4):
            for c in row:
                c.font, c.alignment, c.border = BODY_FONT, WRAP, BORDER
        style_sheet(ws4, [24, 40, 40, 14])

    wb.save(path)
    return path


# ================================================================= 补充资料清单
def build_supplement(data, path):
    meta = data.get("meta", {})
    items = data.get("补充资料清单")
    if not items:
        items = []
        for it in data.get("点检表") or []:
            if it.get("结论") in ("不符合", "待补"):
                items.append({
                    "事项": it.get("说明") or it.get("点检项", ""),
                    "对应点检项": f'{it.get("编号","")} {it.get("点检项","")}',
                    "用途说明": f'点检结论：{it.get("结论")}；证据出处：{it.get("证据出处") or "未找到"}',
                    "紧急度": "上会前" if it.get("结论") == "不符合" else "放款前",
                })
    lines = [
        f'# {meta.get("客户名称","")} 授信审查 · 补充资料清单',
        "",
        f'- 授信金额：{fmt(num(meta.get("授信金额")))} 万元　品种：{meta.get("品种","")}　期限：{meta.get("期限","")}',
        f'- 申报部门：{meta.get("申报部门","")}　主办客户经理：{meta.get("主办客户经理","")}',
        f'- 审查员：{meta.get("审查员","")}　审查日期：{meta.get("审查日期", date.today().isoformat())}',
        f'- 审查结论：{meta.get("审查结论","")}',
        "",
        f'> 共 {len(items)} 项待补。**"上会前"项目未闭环的，不予提交贷审会；'
        f'"放款前"项目未落实的，不得办理首笔提款。**',
        "",
        "| 序号 | 补充事项 | 对应点检项 | 用途说明 | 紧急度 |",
        "|---|---|---|---|---|",
    ]
    order = {"上会前": 0, "放款前": 1, "放款后": 2}
    items = sorted(items, key=lambda x: order.get(x.get("紧急度", "放款前"), 3))
    for i, it in enumerate(items, 1):
        cell = lambda s: str(s or "").replace("\n", " ").replace("|", "／")
        lines.append(f'| {i} | {cell(it.get("事项"))} | {cell(it.get("对应点检项"))} | '
                     f'{cell(it.get("用途说明"))} | {cell(it.get("紧急度"))} |')
    if not items:
        lines.append("| — | 无待补事项 | — | — | — |")
    lines += ["", "---", "",
              "**说明：** 本清单由《审查点检表》中判定为「不符合」「待补」的条目自动汇总生成，",
              "请业务团队逐项落实后回传，并在回传件中标注对应序号。"]
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return path


# ================================================================= 额度测算底稿
def _lab(ws, r, label, value=None, note="", bold=False):
    ws.cell(row=r, column=1, value=label).font = Font(name="微软雅黑", size=10, bold=bold)
    if value is not None:
        c = ws.cell(row=r, column=2, value=value)
        c.font = Font(name="微软雅黑", size=10, bold=bold)
        c.number_format = "#,##0.00"
    if note:
        ws.cell(row=r, column=3, value=note).font = Font(name="微软雅黑", size=9, color="808080")
    return r + 1


def build_worksheet(data, calcs, path):
    wb = openpyxl.Workbook()
    lim_in = (data.get("额度测算") or {}).get("营运资金缺口法") or {}

    # ---------- 营运资金缺口法（活公式）
    ws = wb.active
    ws.title = "营运资金缺口法"
    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 18
    ws.column_dimensions["C"].width = 58
    ws["A1"] = "营运资金缺口法测算底稿（单位：万元）"
    ws["A1"].font = TITLE_FONT
    ws["A2"] = "灰色为输入参数，可直接改；其余为 Excel 公式，改参数后自动重算"
    ws["A2"].font = Font(name="微软雅黑", size=9, color="808080")

    r = 4
    ws.cell(row=r, column=1, value="一、输入参数").font = Font(name="微软雅黑", size=11, bold=True)
    r += 1
    in_start = r
    fields = ["上年度销售收入", "上年度净利润", "营业成本", "预计销售收入年增长率",
              "存货_期初", "存货_期末", "应收账款_期初", "应收账款_期末",
              "预付账款_期初", "预付账款_期末", "应付账款_期初", "应付账款_期末",
              "预收账款_期初", "预收账款_期末",
              "自有资金", "现有流动资金贷款", "其他渠道营运资金",
              "申请额度", "报告测算缺口"]
    ref = {}
    for f in fields:
        v = lim_in.get(f)
        ws.cell(row=r, column=1, value=f.replace("_", " ")).font = BODY_FONT
        if f == "预计销售收入年增长率":
            v = num(v)  # 缺省视为 0，不计入"参数缺失"
        c = ws.cell(row=r, column=2, value=(num(v) if has(v) else None))
        c.font = BODY_FONT
        c.number_format = "0.00%" if f == "预计销售收入年增长率" else "#,##0.00"
        c.fill = PatternFill("solid", fgColor="F2F2F2")
        if not has(v) and f != "预计销售收入年增长率":
            ws.cell(row=r, column=3, value="⚠ 参数缺失 —— 列入待补").font = Font(
                name="微软雅黑", size=9, color="C00000")
        ref[f] = f"B{r}"
        r += 1
    ws.cell(row=r, column=1, value="销售利润率口径").font = BODY_FONT
    cal = ws.cell(row=r, column=2, value=(lim_in.get("销售利润率口径") or "净利率"))
    cal.font, cal.fill = BODY_FONT, PatternFill("solid", fgColor="F2F2F2")
    ws.cell(row=r, column=3, value='填「净利率」或「毛利率」—— 必须与业务团队口径一致，否则差异全部来自口径').font = Font(
        name="微软雅黑", size=9, color="C00000")
    ref["销售利润率口径"] = f"B{r}"
    r += 1
    in_end = r - 1

    r += 1
    ws.cell(row=r, column=1, value="二、复算过程").font = Font(name="微软雅黑", size=11, bold=True)
    r += 1
    steps = [
        ("平均存货", f'=({ref["存货_期初"]}+{ref["存货_期末"]})/2', "(期初+期末)/2"),
        ("平均应收账款", f'=({ref["应收账款_期初"]}+{ref["应收账款_期末"]})/2', ""),
        ("平均预付账款", f'=({ref["预付账款_期初"]}+{ref["预付账款_期末"]})/2', ""),
        ("平均应付账款", f'=({ref["应付账款_期初"]}+{ref["应付账款_期末"]})/2', ""),
        ("平均预收账款", f'=({ref["预收账款_期初"]}+{ref["预收账款_期末"]})/2', ""),
    ]
    base = r
    for i, (n_, f_, note) in enumerate(steps):
        ws.cell(row=r, column=1, value=n_).font = BODY_FONT
        c = ws.cell(row=r, column=2, value=f_)
        c.font, c.number_format = BODY_FONT, "#,##0.00"
        ws.cell(row=r, column=3, value=note).font = Font(name="微软雅黑", size=9, color="808080")
        r += 1
    avg = {n_: f"B{base+i}" for i, (n_, _, _) in enumerate(steps)}

    steps2 = [
        ("存货周转天数", f'=IF({ref["营业成本"]}=0,"",360*{avg["平均存货"]}/{ref["营业成本"]})', "360×平均存货÷营业成本"),
        ("应收账款周转天数", f'=IF({ref["上年度销售收入"]}=0,"",360*{avg["平均应收账款"]}/{ref["上年度销售收入"]})', "360×平均应收÷营业收入"),
        ("预付账款周转天数", f'=IF({ref["营业成本"]}=0,"",360*{avg["平均预付账款"]}/{ref["营业成本"]})', "360×平均预付÷营业成本"),
        ("应付账款周转天数", f'=IF({ref["营业成本"]}=0,"",360*{avg["平均应付账款"]}/{ref["营业成本"]})', "360×平均应付÷营业成本"),
        ("预收账款周转天数", f'=IF({ref["上年度销售收入"]}=0,"",360*{avg["平均预收账款"]}/{ref["上年度销售收入"]})', "360×平均预收÷营业收入"),
    ]
    base2 = r
    for n_, f_, note in steps2:
        ws.cell(row=r, column=1, value=n_).font = BODY_FONT
        c = ws.cell(row=r, column=2, value=f_)
        c.font, c.number_format = BODY_FONT, "#,##0.00"
        ws.cell(row=r, column=3, value=note).font = Font(name="微软雅黑", size=9, color="808080")
        r += 1
    d = {n_: f"B{base2+i}" for i, (n_, _, _) in enumerate(steps2)}

    r_days = r
    ws.cell(row=r, column=1, value="周转天数合计").font = BODY_FONT
    ws.cell(row=r, column=2, value=f'=N({d["存货周转天数"]})+N({d["应收账款周转天数"]})'
                                   f'+N({d["预付账款周转天数"]})-N({d["应付账款周转天数"]})'
                                   f'-N({d["预收账款周转天数"]})').number_format = "#,##0.00"
    ws.cell(row=r, column=3, value="存货+应收+预付−应付−预收").font = Font(name="微软雅黑", size=9, color="808080")
    r += 1
    r_turn = r
    ws.cell(row=r, column=1, value="营运资金周转次数").font = BODY_FONT
    ws.cell(row=r, column=2, value=f'=IF(B{r_days}<=0,"测算失效",360/B{r_days})').number_format = "#,##0.00"
    ws.cell(row=r, column=3, value="360 ÷ 周转天数合计；≤0 表示无营运资金缺口，法失效").font = Font(name="微软雅黑", size=9, color="808080")
    r += 1
    r_margin = r
    ws.cell(row=r, column=1, value="销售利润率").font = BODY_FONT
    ws.cell(row=r, column=2, value=(
        f'=IF({ref["上年度销售收入"]}=0,"",IF({ref["销售利润率口径"]}="毛利率",'
        f'({ref["上年度销售收入"]}-{ref["营业成本"]})/{ref["上年度销售收入"]},'
        f'{ref["上年度净利润"]}/{ref["上年度销售收入"]}))')).number_format = "0.00%"
    ws.cell(row=r, column=3, value="按上方口径取数：毛利率=(收入−成本)/收入；净利率=净利润/收入").font = Font(
        name="微软雅黑", size=9, color="808080")
    r += 1
    r_wc = r
    ws.cell(row=r, column=1, value="营运资金量").font = Font(name="微软雅黑", size=10, bold=True)
    ws.cell(row=r, column=2, value=f'=IF(ISNUMBER(B{r_turn}),{ref["上年度销售收入"]}*(1-B{r_margin})*(1+N({ref["预计销售收入年增长率"]}))/B{r_turn},"测算失效")').number_format = "#,##0.00"
    ws.cell(row=r, column=3, value="销售收入 × (1 − 销售利润率) × (1 + 预计增长率) ÷ 周转次数").font = Font(name="微软雅黑", size=9, color="808080")
    r += 1
    r_gap = r
    ws.cell(row=r, column=1, value="复算新增流贷额度").font = Font(name="微软雅黑", size=10, bold=True)
    ws.cell(row=r, column=2, value=f'=IF(ISNUMBER(B{r_wc}),B{r_wc}-{ref["自有资金"]}-{ref["现有流动资金贷款"]}-{ref["其他渠道营运资金"]},"测算失效")').number_format = "#,##0.00"
    ws.cell(row=r, column=3, value="营运资金量 − 自有资金 − 现有流贷 − 其他渠道").font = Font(name="微软雅黑", size=9, color="808080")
    r += 2

    ws.cell(row=r, column=1, value="三、与报告值比对").font = Font(name="微软雅黑", size=11, bold=True)
    r += 1
    r_diff = r + 1
    r = _lab(ws, r, "报告测算缺口", None)
    ws.cell(row=r - 1, column=2, value=f'={ref["报告测算缺口"]}').number_format = "#,##0.00"
    ws.cell(row=r, column=1, value="差异（复算 − 报告）").font = BODY_FONT
    ws.cell(row=r, column=2, value=f'=IF(ISNUMBER(B{r_gap}),B{r_gap}-B{r_diff-1},"")').number_format = "#,##0.00"
    r += 1
    r_rate = r
    ws.cell(row=r, column=1, value="差异率").font = BODY_FONT
    ws.cell(row=r, column=2, value=f'=IF(OR(B{r_diff-1}=0,NOT(ISNUMBER(B{r_gap}))),"",(B{r_gap}-B{r_diff-1})/B{r_diff-1})').number_format = "0.00%"
    r += 1
    # 关键周转参数留空时，复算必然高估营运资金量 —— 只能判"待补"，不得据此判"不符合"
    blanks = f'COUNTBLANK({ref["应付账款_期初"]}:{ref["预收账款_期末"]})'
    ws.cell(row=r, column=1, value="判定").font = Font(name="微软雅黑", size=10, bold=True)
    ws.cell(row=r, column=2, value=(
        f'=IF({blanks}>0,"待补：应付/预收账款参数缺失，复算高估营运资金量，须补明细账龄表后重算",'
        f'IF(NOT(ISNUMBER(B{r_gap})),"待补：复算失效",'
        f'IF(B{r_rate}="","待补：报告未给测算值",'
        f'IF(ABS(B{r_rate})<=0.1,"符合（差异≤10%）","不符合（差异>10%）"))))'
    )).font = Font(name="微软雅黑", size=10, bold=True)
    r += 1
    ws.cell(row=r, column=1, value="参数完整性").font = BODY_FONT
    ws.cell(row=r, column=2, value=f'=IF(COUNTBLANK({ref["上年度销售收入"]}:{ref["报告测算缺口"]})=0,"参数完整",'
                                   f'"缺失 "&COUNTBLANK({ref["上年度销售收入"]}:{ref["报告测算缺口"]})&" 项，见上方红字标注")').font = BODY_FONT
    r += 2

    ws.cell(row=r, column=1, value="四、额度建议").font = Font(name="微软雅黑", size=11, bold=True)
    r += 1
    ws.cell(row=r, column=1, value="申请额度").font = BODY_FONT
    ws.cell(row=r, column=2, value=f'={ref["申请额度"]}').number_format = "#,##0.00"
    r += 1
    ws.cell(row=r, column=1, value="路径 A：按复算缺口核定").font = BODY_FONT
    ws.cell(row=r, column=2, value=f'=IF(ISNUMBER(B{r_gap}),MIN({ref["申请额度"]},MAX(B{r_gap},0)),"—")').number_format = "#,##0.00"
    r += 1
    ws.cell(row=r, column=1, value="路径 B：转订单模式").font = BODY_FONT
    ws.cell(row=r, column=2, value="见「订单模式」工作表").font = BODY_FONT
    r += 1
    ws.cell(row=r, column=1, value="※ 由审查员决策，本表只并列两套结果，不自动写死结论").font = Font(
        name="微软雅黑", size=9, color="C00000")

    # ---------- 订单模式
    ws2 = wb.create_sheet("订单模式")
    for col, w in zip("ABCDEF", (10, 26, 20, 18, 12, 18)):
        ws2.column_dimensions[col].width = w
    ws2["A1"] = "订单模式测算底稿（营运资金缺口法失效或缺口小于申请额时使用）"
    ws2["A1"].font = TITLE_FONT
    ws2.append([])
    ws2.append(["序号", "供应商", "标的", "合同金额（万元）", "融资比例", "可融资额（万元）"])
    hdr = ws2.max_row
    od_in = (data.get("额度测算") or {}).get("订单模式") or {}
    orders = od_in.get("订单") or [{"供应商": "", "标的": "", "金额": None, "融资比例": 0.7}]
    first = hdr + 1
    for i, o in enumerate(orders, 1):
        ws2.append([i, o.get("供应商", ""), o.get("标的", ""),
                    (num(o.get("金额")) if has(o.get("金额")) else None),
                    num(o.get("融资比例"), 0.7), None])
        ws2.cell(row=ws2.max_row, column=6, value=f"=D{ws2.max_row}*E{ws2.max_row}")
    last = ws2.max_row
    ws2.append(["合计", "", "", f"=SUM(D{first}:D{last})", "", f"=SUM(F{first}:F{last})"])
    tot_row = ws2.max_row
    for row in ws2.iter_rows(min_row=hdr, max_row=tot_row, max_col=6):
        for c in row:
            c.font, c.border = BODY_FONT, BORDER
            if c.column in (4, 6):
                c.number_format = "#,##0.00"
            if c.column == 5:
                c.number_format = "0%"
    for c in ws2[hdr]:
        c.fill, c.font, c.alignment = HEAD_FILL, HEAD_FONT, CENTER

    r = tot_row + 2
    ws2.cell(row=r, column=1, value="期限倒推（天）").font = Font(name="微软雅黑", size=11, bold=True)
    r += 1
    day_rows = {}
    for k in ("采购交付天数", "施工验收天数", "结算天数", "账期天数", "缓冲天数"):
        ws2.cell(row=r, column=1, value=k).font = BODY_FONT
        c = ws2.cell(row=r, column=2, value=(num(od_in.get(k)) if has(od_in.get(k)) else None))
        c.font, c.fill = BODY_FONT, PatternFill("solid", fgColor="F2F2F2")
        day_rows[k] = r
        r += 1
    ws2.cell(row=r, column=1, value="回款周期合计").font = Font(name="微软雅黑", size=10, bold=True)
    ws2.cell(row=r, column=2, value=f'=SUM(B{min(day_rows.values())}:B{max(day_rows.values())})')
    tot_days = r
    r += 1
    ws2.cell(row=r, column=1, value="建议授信期限（月）").font = Font(name="微软雅黑", size=10, bold=True)
    ws2.cell(row=r, column=2, value=f'=IF(B{tot_days}=0,"—",ROUNDUP(B{tot_days}/30,0))')
    r += 2
    ws2.cell(row=r, column=1, value="※ 期限不得超过单笔订单的实际回款周期；多笔订单按加权平均确定并以最长单笔封顶。").font = Font(
        name="微软雅黑", size=9, color="C00000")
    r += 1
    ws2.cell(row=r, column=1, value="※ 额度上限同时受：① 申请额度 ② 单一客户 5000 万元限额 ③ 合同金额总和 三者约束。").font = Font(
        name="微软雅黑", size=9, color="C00000")

    # ---------- 担保能力
    g_in = data.get("担保能力") or {}
    ws3 = wb.create_sheet("担保能力")
    for col, w in zip("ABC", (28, 18, 62)):
        ws3.column_dimensions[col].width = w
    ws3["A1"] = "担保能力测算底稿（单位：万元）"
    ws3["A1"].font = TITLE_FONT
    r = 3
    ws3.cell(row=r, column=1, value="担保人").font = BODY_FONT
    ws3.cell(row=r, column=2, value=g_in.get("担保人", "")).font = BODY_FONT
    r += 1
    gref = {}
    for k in ("净资产", "已担保金额", "或有负债", "受限资产", "本次授信金额", "关注类担保余额"):
        ws3.cell(row=r, column=1, value=k).font = BODY_FONT
        c = ws3.cell(row=r, column=2, value=(num(g_in.get(k)) if has(g_in.get(k)) else None))
        c.font, c.number_format = BODY_FONT, "#,##0.00"
        c.fill = PatternFill("solid", fgColor="F2F2F2")
        gref[k] = f"B{r}"
        r += 1
    ws3.cell(row=r, column=1, value="担保能力").font = Font(name="微软雅黑", size=10, bold=True)
    ws3.cell(row=r, column=2, value=f'={gref["净资产"]}-{gref["已担保金额"]}-{gref["或有负债"]}-{gref["受限资产"]}').number_format = "#,##0.00"
    ws3.cell(row=r, column=3, value="净资产 − 已担保 − 或有负债 − 受限资产").font = Font(name="微软雅黑", size=9, color="808080")
    r_cap = r
    r += 1
    ws3.cell(row=r, column=1, value="覆盖倍数").font = BODY_FONT
    ws3.cell(row=r, column=2, value=f'=IF({gref["本次授信金额"]}=0,"",B{r_cap}/{gref["本次授信金额"]})').number_format = "#,##0.00"
    r += 1
    ws3.cell(row=r, column=1, value="判定").font = Font(name="微软雅黑", size=10, bold=True)
    ws3.cell(row=r, column=2, value=f'=IF(B{r_cap}>={gref["本次授信金额"]},"符合","不符合")').font = Font(name="微软雅黑", size=10, bold=True)
    r += 2
    for label, key in (("净资产口径", "净资产口径"), ("已担保口径", "已担保口径"), ("基准日", "基准日")):
        ws3.cell(row=r, column=1, value=label).font = BODY_FONT
        ws3.cell(row=r, column=2, value=g_in.get(key, "⚠ 未注明")).font = BODY_FONT
        r += 1
    ws3.cell(row=r, column=1, value="※ 净资产须统一口径（优先归母），已担保优先取责任金额而非余额；"
                                    "关注类担保应全额扣减并单独评估代偿风险。").font = Font(
        name="微软雅黑", size=9, color="C00000")

    # ---------- 财务勾稽
    ties = calcs.get("勾稽")
    if ties:
        ws4 = wb.create_sheet("财务勾稽")
        for col, w in zip("ABCDEFG", (8, 10, 32, 16, 16, 14, 46)):
            ws4.column_dimensions[col].width = w
        ws4.append(["编号", "年份", "校验项", "计算值", "报告值", "差异", "判定"])
        for rec in ties:
            for c in rec["检查"]:
                ws4.append([c["编号"], rec["年份"], c["项目"], c["计算值"],
                            c["报告值"], c["差异"], c["判定"]])
        for row in ws4.iter_rows(min_row=2, max_row=ws4.max_row, max_col=7):
            for c in row:
                c.font, c.alignment, c.border = BODY_FONT, WRAP, BORDER
            v = str(row[6].value or "")
            if any(k in v for k in ("不一致", "不平衡", "存疑", "偏大", "不成立", "弱")):
                row[6].fill = FILL["不符合"]
            elif "待补" in v:
                row[6].fill = FILL["待补"]
        style_sheet(ws4, [8, 10, 32, 16, 16, 14, 46])

    wb.save(path)
    return path


# ================================================================= main
def main():
    ap = argparse.ArgumentParser(description="审查四交付物生成器")
    ap.add_argument("review_json")
    ap.add_argument("-o", "--outdir", default=".")
    args = ap.parse_args()

    with open(args.review_json, encoding="utf-8") as f:
        data = json.load(f)
    os.makedirs(args.outdir, exist_ok=True)

    meta = data.get("meta", {})
    name = meta.get("客户名称", "客户")
    amt = f'{num(meta.get("授信金额")):.0f}'
    prefix = f"{name}-{amt}万元"

    lim_in = (data.get("额度测算") or {}).get("营运资金缺口法")
    lim = calc_limit(lim_in)
    # 融资必要性以「业务部门申报的资金缺口」为分母；缺失时退回复算值
    gap = num((lim_in or {}).get("报告测算缺口")) or None
    if not gap and lim and lim.get("ok"):
        gap = lim.get("复算缺口")
    calcs = {
        "额度": lim,
        "订单": calc_order((data.get("额度测算") or {}).get("订单模式")),
        "担保": calc_guarantee(data.get("担保能力")),
        "勾稽": calc_ties(data.get("财务勾稽"), gap=gap),
    }

    made = []
    made.append(build_docx(data, calcs, os.path.join(args.outdir, f"{prefix}-审查报告.docx")))
    made.append(build_checklist(data, os.path.join(args.outdir, f"{prefix}-审查点检表.xlsx")))
    made.append(build_supplement(data, os.path.join(args.outdir, f"{prefix}-补充资料清单.md")))
    made.append(build_worksheet(data, calcs, os.path.join(args.outdir, f"{prefix}-额度测算底稿.xlsx")))

    items = data.get("点检表") or []
    n_bad = sum(1 for i in items if i.get("结论") == "不符合")
    n_pend = sum(1 for i in items if i.get("结论") == "待补")
    for p in made:
        print("✓", p)
    print(f"\n点检 {len(items)} 项：不符合 {n_bad} 项，待补 {n_pend} 项。")
    if n_bad and meta.get("审查结论") == "拟同意":
        print("⚠ 存在「不符合」项却出具「拟同意」—— 违反 SKILL 阶段 4 规则，请复核审查结论。")
    if len(data.get("穿透抽查") or []) < 2:
        print("⚠ 穿透抽查不足 2 项 —— 违反反套路化要求，请返回阶段 3 补做。")


if __name__ == "__main__":
    sys.exit(main())
